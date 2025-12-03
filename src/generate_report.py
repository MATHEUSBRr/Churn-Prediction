import os
from docx import Document
from docx.shared import Inches
BASE = os.path.dirname(os.path.dirname(__file__))
DOC_OUT = os.path.join(BASE, "reports", "report.docx")
FIG_DIR = os.path.join(BASE, "reports", "figures")
METRICS = os.path.join(BASE, "reports", "metrics.txt")

def build_report():
    doc = Document()
    doc.add_heading('Projeto PP - Previsão de Churn', level=1)
    doc.add_paragraph('Resumo: Pipeline completo (ETL, EDA, modelagem, avaliação).')
    doc.add_heading('1. Dados', level=2)
    doc.add_paragraph('Dataset sintético com variáveis demográficas, uso e cobrança.')
    doc.add_heading('2. Metodologia', level=2)
    doc.add_paragraph('Foi realizado tratamento de dados, engenharia de features, treino de RandomForest e avaliação com métricas clássicas.')
    doc.add_heading('3. Métricas', level=2)
    if os.path.exists(METRICS):
        with open(METRICS,'r') as f:
            doc.add_paragraph(f.read())
    doc.add_heading('4. Figuras', level=2)
    for fname in ['churn_dist.png','age_hist.png','tenure_vs_monthly.png','roc.png','confusion.png','feat_imp.png']:
        p = os.path.join(FIG_DIR, fname)
        if os.path.exists(p):
            doc.add_paragraph(fname)
            try:
                doc.add_picture(p, width=Inches(5))
            except Exception as e:
                doc.add_paragraph("Figura não inserida: " + str(e))
    doc.add_heading('Conclusão', level=2)
    doc.add_paragraph('Interpretação: features com maior impacto podem ser usadas em campanhas de retenção.')
    doc.save(DOC_OUT)
    print('Report generated at', DOC_OUT)

if __name__ == '__main__':
    build_report()
